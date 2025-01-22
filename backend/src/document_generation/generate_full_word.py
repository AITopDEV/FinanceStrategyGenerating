from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from docx.shared import Pt
from docx.shared import RGBColor
import json


def generate_title(doc, client_names, image_path, result_path):
    names_combined = " and ".join(client_names)
    # Replace placeholder with the combined names
    placeholder = "Client Names"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder, names_combined)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(14)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)
    if image_path:
        doc.add_picture(image_path)
    else:
        doc.add_picture("temp_word/title.png")
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/1.title.docx")
    return doc


def generate_content(doc, result_path):
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/2.contents.docx")
    return doc


def generate_requirements_objectives(doc, needs, requirements, objectives, reason, result_path):
    # Replace placeholder with the combined names
    placeholder_needs = "«needs»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_needs in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_needs, needs)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)

    # Replace placeholder with the combined names
    placeholder_requirements = "«requirements»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_requirements in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_requirements, requirements)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)

    # Replace placeholder with the combined names
    placeholder_objectives = "«objectives»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_objectives in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_objectives, objectives)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)

    # Replace placeholder with the combined names
    placeholder_reason = "«reason»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_reason in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_reason, reason)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/3.requirements_objectives.docx")
    return doc


def generate_product_selection(doc, result_path):
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/4.product_selection.docx")
    return doc


def generate_highlight_solution(doc, proposed_credit_application, result_path):
    # Replace placeholder with the combined names
    placeholder_name = "«clientNames»"
    placeholder_lender = "«lender»"
    placeholder_product_name = "«productName»"
    placeholder_product_amount = "«productAmount»"
    placeholder_interest_rate = "«rateCard»"
    placeholder_repayment = "«repaymentCard»"
    placeholder_loan_term = "«loanTerm»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_name in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_name, proposed_credit_application["application name"])

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = True  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)  

        if placeholder_lender in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_lender, proposed_credit_application["credit provider"])

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = True  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)  
            # Replace placeholders in tables (if needed)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    if placeholder_product_name in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_product_name,
                                                      proposed_credit_application["product name"])

                        # Clear all the existing runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ""

                            # Add a single new run with the updated text
                        new_run = paragraph.add_run(full_text)

                        # Set font attributes for the entire replaced text
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(12)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    if placeholder_product_amount in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_product_amount,
                                                      proposed_credit_application["loan amount"])

                        # Clear all the existing runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ""

                            # Add a single new run with the updated text
                        new_run = paragraph.add_run(full_text)

                        # Set font attributes for the entire replaced text
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(12)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    if placeholder_interest_rate in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_interest_rate,
                                                      proposed_credit_application["interest rate"])

                        # Clear all the existing runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ""

                            # Add a single new run with the updated text
                        new_run = paragraph.add_run(full_text)

                        # Set font attributes for the entire replaced text
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(12)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    if placeholder_repayment in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_repayment, proposed_credit_application["repayments"])

                        # Clear all the existing runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ""

                            # Add a single new run with the updated text
                        new_run = paragraph.add_run(full_text)

                        # Set font attributes for the entire replaced text
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(12)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    if placeholder_loan_term in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_loan_term, proposed_credit_application["loan term"])

                        # Clear all the existing runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ""

                            # Add a single new run with the updated text
                        new_run = paragraph.add_run(full_text)

                        # Set font attributes for the entire replaced text
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(12)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/5.highlight_solution.docx")
    return doc


def generate_product_comparison(doc, names, product_comparison, result_path):
    # Replace placeholder with the combined names
    placeholder_name = "«clientNames»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_name in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_name, names)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)  

    transformed_data = {
        "Product providers": [],
        "Product names": [],
        "Loan amounts": [],
        "Interest rates": [],
        "Repayments": [],
        "Features": [],
        "App fees": [],
        "Monthly fees": [],
        "Annual fees": [],
        "Discharge fees": [],
        "Total interest": [],
        "Total monthly fees": [],
        "Total annual fees": [],
        "Total setup costs": [],
        "Loan terms": [],
        "Total loan costs": [],
        "Savings": [],
        "Comparative savings": []
    }

    # Loop through each product and append to the new structure
    for product in product_comparison:
        transformed_data["Product providers"].append(product.get("product provider", "N/A"))
        transformed_data["Product names"].append(product.get("product name", "N/A"))
        transformed_data["Loan amounts"].append(product.get("loan amount", "N/A"))
        transformed_data["Interest rates"].append(product.get("interest rates", "N/A"))
        transformed_data["Repayments"].append(product.get("repayments", "N/A"))
        transformed_data["Features"].append(product.get("features", "N/A"))
        transformed_data["App fees"].append(product.get("app fee", "N/A"))
        transformed_data["Monthly fees"].append(product.get("monthly fee", "N/A"))
        transformed_data["Annual fees"].append(product.get("annual fee", "N/A"))
        transformed_data["Discharge fees"].append(product.get("discharge fees", "N/A"))
        transformed_data["Total interest"].append(product.get("total interest", "N/A"))
        transformed_data["Total monthly fees"].append(product.get("total monthly fees", "N/A"))
        transformed_data["Total annual fees"].append(product.get("total of annual fees", "N/A"))
        transformed_data["Total setup costs"].append(product.get("total setup costs", "N/A"))
        transformed_data["Loan terms"].append(product.get("loan term", "N/A"))
        transformed_data["Total loan costs"].append(product.get("total loan cost", "N/A"))
        transformed_data["Savings"].append(product.get("saving", "N/A"))
        transformed_data["Comparative savings"].append(product.get("comparative saving", "N/A"))

    # Transpose data: Fields are rows, providers are columns
    fields = list(transformed_data.keys())  # Rows: Field names
    providers = transformed_data["Product providers"]  # Columns: Product providers

    table = doc.add_table(rows=len(fields), cols=len(providers) + 1)
    styles = doc.styles
    print(styles)

    # Populate the data rows (fields as rows, values as columns)
    for row_idx, field in enumerate(fields):
        cell = table.rows[row_idx].cells[0]
        # First cell in each row is the field name
        table.rows[row_idx].cells[0].text = field.replace("_", " ").capitalize()
        # Apply line spacing to the text inside the field name cell
        for paragraph in cell.paragraphs:  # Each cell can have multiple paragraphs
            paragraph.paragraph_format.line_spacing = Pt(12)  # Set spacing (adjust as needed)
            paragraph.paragraph_format.space_after = Pt(6)  # Space after paragraph
            paragraph.paragraph_format.space_before = Pt(6)  # Space before paragraph

        # Remaining cells in the row are the values for the providers
        for col_idx, value in enumerate(transformed_data[field]):
            cell = table.rows[row_idx].cells[col_idx + 1]
            cell.text = str(value)

            # Apply line spacing to the text inside each provider's cell
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = Pt(12)  # Set spacing
                paragraph.paragraph_format.space_after = Pt(6)  # Space after paragraph
                paragraph.paragraph_format.space_before = Pt(6)  # Space before paragraph

    table.style = styles["List Table 1 Light Accent 1"]
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/6.product_comparison.docx")
    return doc


def generate_funding_position(doc, names, funding_position, result_path):
    # Replace placeholder with the combined names
    placeholder_name = "«clientNames»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_name in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_name, names)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)  

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    placeholder_purchase_amount = "«purchase_amount»"
                    if placeholder_purchase_amount in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_purchase_amount,
                                                      funding_position["funds Required"]["purchase amount"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_refinance_amount = "«refinance_amount»"
                    if placeholder_refinance_amount in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_refinance_amount,
                                                      funding_position["funds Required"]["refinance amount"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_purchase_sd = "«purchase_sd»"
                    if placeholder_purchase_sd in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_purchase_sd,
                                                      funding_position["funds Required"][
                                                          "stamp duty - transfer of land"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_registration_fee = "«registration_fee»"
                    if placeholder_registration_fee in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_registration_fee,
                                                      funding_position["funds Required"][
                                                          "titles office - mortgage registration"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_mortgage_discharge_fee = "«mortgage_discharge_fee»"
                    if placeholder_mortgage_discharge_fee in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_mortgage_discharge_fee,
                                                      funding_position["funds Required"][
                                                          "titles office - mortgage discharge"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_transfer_fee = "«transfer_fee»"
                    if placeholder_transfer_fee in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_transfer_fee,
                                                      funding_position["funds Required"][
                                                          "titles office - transfer of land"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_establishment_fee = "«establishment_fee»"
                    if placeholder_establishment_fee in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_establishment_fee,
                                                      funding_position["funds Required"][
                                                          "establishment fee"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_legal_fee = "«legal_fee»"
                    if placeholder_legal_fee in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_legal_fee,
                                                      funding_position["funds Required"][
                                                          "legal costs"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_foreign_buyer_surcharge = "«foreign_buyer_surcharge»"
                    if placeholder_foreign_buyer_surcharge in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_foreign_buyer_surcharge,
                                                      funding_position["funds Required"][
                                                          "foreign buyer fees"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_discharge_cost = "«discharge_cost»"
                    if placeholder_discharge_cost in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_discharge_cost,
                                                      funding_position["funds Required"][
                                                          "discharge cost"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_lmi = "«lmi»"
                    if placeholder_lmi in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_lmi,
                                                      funding_position["funds Required"][
                                                          "lenders mortgage insurance"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_fee_other = "«fee_other»"
                    if placeholder_fee_other in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_fee_other,
                                                      funding_position["funds Required"][
                                                          "other/sundries"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_loan_requested = "«loan_requested»"
                    if placeholder_loan_requested in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_loan_requested,
                                                      funding_position["funds available"][
                                                          "loan sought"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_fhog = "«fhog»"
                    if placeholder_fhog in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_fhog,
                                                      funding_position["funds available"][
                                                          "FHOG"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_grossSaleProceeds = "«grossSaleProceeds»"
                    if placeholder_grossSaleProceeds in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_grossSaleProceeds,
                                                      funding_position["funds available"][
                                                          "sale proceeds"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_funds = "«funds»"
                    if placeholder_funds in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_funds,
                                                      funding_position["funds available"][
                                                          "other funds available"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_debts = "«debts»"
                    if placeholder_debts in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_debts,
                                                      funding_position["funds available"][
                                                          "debts to repay"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_funds_subtotal = "«funds_subtotal»"
                    if placeholder_funds_subtotal in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_funds_subtotal,
                                                      funding_position["total"][
                                                          "sub total funds required"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_deposit_paid = "«deposit_paid»"
                    if placeholder_deposit_paid in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_deposit_paid,
                                                      funding_position["total"][
                                                          "deposit already paid"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_funds_required = "«funds_required»"
                    if placeholder_funds_required in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_funds_required,
                                                      funding_position["total"][
                                                          "total funds required"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_total_lend = "«total_lend»"
                    if placeholder_total_lend in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_total_lend,
                                                      funding_position["total"][
                                                          "total lend"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_total_security = "«total_security»"
                    if placeholder_total_security in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_total_security,
                                                      funding_position["total"][
                                                          "total security"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_lvr = "«lvr»"
                    if placeholder_lvr in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_lvr,
                                                      funding_position["total"][
                                                          "loan value ratio"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_funds_available = "«funds_available»"
                    if placeholder_funds_available in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_funds_available,
                                                      funding_position["total"][
                                                          "total funds available"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)  

                    placeholder_deficit_surplus_amount = "«deficit_surplus_amount»"
                    if placeholder_deficit_surplus_amount in full_text:
                        # Replace the placeholder text
                        full_text = full_text.replace(placeholder_deficit_surplus_amount,
                                                      funding_position["funds surplus"])
                        for run in paragraph.runs:
                            run.text = ""
                        new_run = paragraph.add_run(full_text)
                        new_run.font.name = 'Nunito Sans'  # Font type
                        new_run.font.size = Pt(8)  # Font size
                        new_run.font.bold = False  # Bold text
                        new_run.font.italic = False  # Italicize if needed
                        new_run.font.color.rgb = RGBColor(0, 0, 0)
    # Add a page break
    paragraph = doc.add_paragraph()  # Create a new paragraph
    run = paragraph.add_run()  # Add a run (a container for text)
    run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/7.funding_position.docx")
    return doc


def generate_guide(doc, names, result_path):
    # Replace placeholder with the combined names
    placeholder_name = "«clientNames»"
    # Update font in paragraphs
    for paragraph in doc.paragraphs:
        # Combine the text in all runs of the paragraph
        full_text = "".join(run.text for run in paragraph.runs)

        # Check if the placeholder is in the full text
        if placeholder_name in full_text:
            # Replace the placeholder text
            full_text = full_text.replace(placeholder_name, names)

            # Clear all the existing runs in the paragraph
            for run in paragraph.runs:
                run.text = ""

                # Add a single new run with the updated text
            new_run = paragraph.add_run(full_text)

            # Set font attributes for the entire replaced text
            new_run.font.name = 'Nunito Sans'  # Font type
            new_run.font.size = Pt(12)  # Font size
            new_run.font.bold = False  # Bold text
            new_run.font.italic = False  # Italicize if needed
            new_run.font.color.rgb = RGBColor(0, 0, 0)
    # # Add a page break
    # paragraph = doc.add_paragraph()  # Create a new paragraph
    # run = paragraph.add_run()  # Add a run (a container for text)
    # run.add_break(WD_BREAK.PAGE)  # Add a page break
    doc.save(result_path + "/8.guide.docx")
    return doc


def generate_last(doc, result_path):
    doc.save(result_path + "/9.last.docx")
    return doc


def merge_docx(input_files, output_file):
    # Create a new document for merging
    # Open the first document as the base document
    merged_document = Document(input_files[0])
    composer = Composer(merged_document)

    # Iterate through the rest of the documents and append them to the base document
    for input_file in input_files[1:]:
        doc = Document(input_file)
        composer.append(doc)

        # Save the merged document
    composer.save(output_file)
    return merged_document


def generate_full_docx(preliminary_assessment, credit_proposal, pic_path, temp_path, result_path):
    doc_title = Document(temp_path + "/1.title.docx")
    generate_title(doc_title, preliminary_assessment["name"], pic_path, result_path)

    doc_content = Document(temp_path + "/2.contents.docx")
    generate_content(doc_content, result_path)

    doc_content = Document(temp_path + "/3.requirements_objectives.docx")
    generate_requirements_objectives(doc_content,
                                     credit_proposal["needs"],
                                     credit_proposal["requirements"],
                                     credit_proposal["objectives"],
                                     credit_proposal["reason"],
                                     result_path)

    doc_content = Document(temp_path + "/4.product_selection.docx")
    generate_product_selection(doc_content, result_path)

    doc_highlight_solution = Document(temp_path + "/5.highlight_solution.docx")
    generate_highlight_solution(doc_highlight_solution, credit_proposal["proposed_credit_application"], result_path)

    doc_product_comparison = Document(temp_path + "/6.product_comparison.docx")
    generate_product_comparison(doc_product_comparison,
                                credit_proposal["proposed_credit_application"]["application name"],
                                preliminary_assessment["product_comparison"],
                                result_path)

    doc_funding_position = Document(temp_path + "/7.funding_position.docx")
    generate_funding_position(doc_funding_position,
                              credit_proposal["proposed_credit_application"]["application name"],
                              preliminary_assessment["security_property"],
                              result_path)

    doc_guide = Document(temp_path + "/8.guide.docx")
    generate_guide(doc_guide, credit_proposal["proposed_credit_application"]["application name"], result_path)

    doc_last = Document(temp_path + "/9.last.docx")
    generate_last(doc_last, result_path)

    input_docs = [result_path + "/1.title.docx",
                  result_path + "/2.contents.docx",
                  result_path + "/3.requirements_objectives.docx",
                  result_path + "/4.product_selection.docx",
                  result_path + "/5.highlight_solution.docx",
                  result_path + "/6.product_comparison.docx",
                  result_path + "/7.funding_position.docx",
                  result_path + "/8.guide.docx",
                  result_path + "/9.last.docx", ]
    merge_docx(input_docs, result_path + "/final.docx")


if __name__ == '__main__':
    with open("data/result_preliminary_assessment.json", "r") as file:
        result_preliminary_assessment = json.load(file)  # Parse the JSON file
    with open("data/result_credit_proposal.json", "r") as file:
        result_credit_proposal = json.load(file)  # Parse the JSON file
    house_pic_path = "../../data/result_word/view_1_heading_0_pitch_0.jpg"

    generate_full_docx(result_preliminary_assessment,
                       result_credit_proposal,
                       house_pic_path,
                       "../../data/temp_word",
                       "../../data/result_word")
