import json

from docx import Document


def process_proposed_credit_app(table1, table2):
    proposed_credit_app_info = {
        "application name": table1.cell(0, 1).text,
        "credit provider": table1.cell(1, 1).text,
        "product": table2.cell(0, 1).text,
        "product name": table2.cell(1, 1).text,
        "loan amount": table2.cell(2, 1).text,
        "interest rate": table2.cell(3, 1).text,
        "repayments": table2.cell(4, 1).text,
        "loan term": table2.cell(5, 1).text,
    }

    return proposed_credit_app_info


def parser_doc(doc_path):
    doc = Document(doc_path)

    doc_data = {}

    name_index = 0
    requirements_objectives_index = 0
    reason_start_index = 0
    reason_end_index = 0

    for para_index, para in enumerate(doc.paragraphs):
        if para.text == "Prepared for:":
            name_index = para_index + 1
        if para.text == "Your requirements and objectives":
            requirements_objectives_index = para_index + 1
        if para.text == "Reason why this product was selected":
            reason_start_index = para_index + 1
        if ("This recommendation aligns well with your objectives and is considered not unsuitable for your needs." in
                para.text):
            reason_end_index = para_index + 1

    names = [name.strip() for name in doc.paragraphs[name_index].text.split('&')]
    doc_data["name"] = names
    doc_data["requirements_objectives"] = doc.paragraphs[requirements_objectives_index].text
    doc_data["reason"] = " ".join(
        doc.paragraphs[i].text.strip() for i in range(reason_start_index, reason_end_index + 1)
        if doc.paragraphs[i].text.strip()
    )
    proposed_credit_app_table_index = 0

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell.text == "Proposed Credit Application":
                    proposed_credit_app_table_index = table_index + 1

    proposed_credit_app = process_proposed_credit_app(doc.tables[proposed_credit_app_table_index],
                                                      doc.tables[proposed_credit_app_table_index + 1])
    doc_data["proposed_credit_application"] = proposed_credit_app

    return doc_data


if __name__ == '__main__':
    filepath = '../../references/Credit Proposal/Delaney & Ralley - Credit Proposal.docx'
    with open("../document_generation/data/result_credit_proposal.json", "w") as json_file:
        json.dump(parser_doc(filepath), json_file, indent=4)
    print(parser_doc(filepath))
