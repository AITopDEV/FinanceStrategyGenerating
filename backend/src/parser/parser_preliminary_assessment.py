import json

from docx import Document


def process_personal_info_table(table):
    personal_info = {
        "name": table.cell(0, 0).text[:-1],
        "address": table.cell(1, 1).text,
        "income": table.cell(2, 1).text,
        "expenses": table.cell(3, 1).text,
    }
    return personal_info


def process_security_property(table1, table2):
    security_property_info = {
        "funds Required": {
            "purchase amount": table1.cell(1, 1).text,
            "refinance amount": table1.cell(2, 1).text,
            "stamp duty - transfer of land": table1.cell(3, 1).text,
            "titles office - mortgage registration": table1.cell(4, 1).text,
            "titles office - mortgage discharge": table1.cell(5, 1).text,
            "titles office - transfer of land": table1.cell(6, 1).text,
            "establishment fee": table1.cell(7, 1).text,
            "legal costs": table1.cell(8, 1).text,
            "foreign buyer fees": table1.cell(9, 1).text,
            "discharge cost": table1.cell(10, 1).text,
            "lenders mortgage insurance": table1.cell(11, 1).text,
            "other/sundries": table1.cell(12, 1).text,

        },
        "funds available": {
            "loan sought": table1.cell(1, 3).text,
            "FHOG": table1.cell(2, 3).text,
            "sale proceeds": table1.cell(3, 3).text,
            "other funds available": table1.cell(5, 2).text,
            "debts to repay": table1.cell(11, 2).text,
        },
        "total": {
            "sub total funds required": table2.cell(0, 1).text,
            "deposit already paid": table2.cell(1, 1).text,
            "total funds required": table2.cell(3, 1).text,
            "total lend": table2.cell(0, 3).text,
            "total security": table2.cell(1, 3).text,
            "loan value ratio": table2.cell(2, 3).text,
            "total funds available": table2.cell(3, 3).text,
        },
        "funds surplus": table2.cell(4, 0).text.split("Funds Surplus", 1)[1].strip(),
    }

    return security_property_info


def process_product_recommendation(table):
    product_recommendation = []
    for row_index, row in enumerate(table.rows):
        each_product_recommendation = {}
        if row_index != 0:
            each_product_recommendation["product provider"] = table.cell(row_index, 0).text
            each_product_recommendation["product name"] = table.cell(row_index, 1).text
            each_product_recommendation["interest rate"] = table.cell(row_index, 2).text
            each_product_recommendation["total setup costs"] = table.cell(row_index, 3).text
            each_product_recommendation["ongoing fees"] = table.cell(row_index, 4).text
            each_product_recommendation["total cost of loan"] = table.cell(row_index, 5).text
            product_recommendation.append(each_product_recommendation)
    return product_recommendation


def process_consumer_requirements_analysis1(table):
    consumer_requirements_analysis1 = []
    for row_index, row in enumerate(table.rows):
        each_consumer_requirements_analysis1 = {}
        if row_index != 0:
            each_consumer_requirements_analysis1["number"] = table.cell(row_index, 0).text
            each_consumer_requirements_analysis1["question"] = table.cell(row_index, 1).text
            each_consumer_requirements_analysis1["response"] = table.cell(row_index, 2).text
            consumer_requirements_analysis1.append(each_consumer_requirements_analysis1)
    return consumer_requirements_analysis1


def process_consumer_requirements_analysis2(table):
    consumer_requirements_analysis2 = []
    for row_index, row in enumerate(table.rows):
        each_consumer_requirements_analysis2 = {}
        if row_index != 0:
            each_consumer_requirements_analysis2["living expenses"] = table.cell(row_index, 0).text
            each_consumer_requirements_analysis2["allocation"] = table.cell(row_index, 1).text
            each_consumer_requirements_analysis2["amount"] = table.cell(row_index, 2).text
            consumer_requirements_analysis2.append(each_consumer_requirements_analysis2)
    return consumer_requirements_analysis2


def process_product_comparison(table, product_recommendation):
    product_comparison = []

    for product in product_recommendation:
        product_comparison.append({
            "product provider": product["product provider"],
            "product name": product["product name"]
        })

    for i in range(1, len(table.rows[1].cells)):
        for product in product_comparison:
            if product["product name"] == table.rows[1].cells[i].text:
                product["loan amount"] = table.cell(2, i).text
                product["interest rates"] = table.cell(3, i).text
                product["repayments"] = table.cell(4, i).text
                product["features"] = table.cell(5, i).text
                product["app fee"] = table.cell(6, i).text
                product["monthly fee"] = table.cell(7, i).text
                product["annual fee"] = table.cell(8, i).text
                product["discharge fees"] = table.cell(9, i).text
                product["total interest"] = table.cell(10, i).text
                product["total monthly fees"] = table.cell(11, i).text
                product["total of annual fees"] = table.cell(12, i).text
                product["total setup costs"] = table.cell(13, i).text
                product["loan term"] = table.cell(14, i).text
                product["total loan cost"] = table.cell(15, i).text
                product["saving"] = table.cell(16, i).text
                product["comparative saving"] = table.cell(17, i).text

    return product_comparison


def parser_doc(doc_path):
    doc = Document(doc_path)

    doc_data = {}

    name_index = 0

    for para_index, para in enumerate(doc.paragraphs):
        if para.text == "Prepared for:":
            name_index = para_index + 1

    names = [name.strip() for name in doc.paragraphs[name_index].text.split('&')]
    doc_data["name"] = names

    personal_info_table_index = 0
    security_property_table_index = 0
    product_recommendation_table_index = 0
    consumer_requirements_analysis_table_index = 0
    product_comparison_types_table_index = 0

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell.text == "Personal Information":
                    personal_info_table_index = table_index + 1
                if "Security Property(s) & Funding Position" in cell.text:
                    security_property_table_index = table_index + 1
                if "Product Recommendation" in cell.text:
                    product_recommendation_table_index = table_index + 1
                    consumer_requirements_analysis_table_index = product_recommendation_table_index + 1
                    product_comparison_types_table_index = consumer_requirements_analysis_table_index + 2

    personal_info = []
    for i in range(personal_info_table_index, security_property_table_index - 1):
        personal_info.append(process_personal_info_table(doc.tables[i]))
    doc_data["personal_info"] = personal_info

    security_property = process_security_property(doc.tables[security_property_table_index],
                                                  doc.tables[security_property_table_index + 1])
    doc_data["security_property"] = security_property

    product_recommendation = process_product_recommendation(doc.tables[product_recommendation_table_index])
    doc_data["product_recommendation"] = product_recommendation

    consumer_requirements_analysis1 = process_consumer_requirements_analysis1(
        doc.tables[consumer_requirements_analysis_table_index])

    consumer_requirements_analysis2 = process_consumer_requirements_analysis2(
        doc.tables[consumer_requirements_analysis_table_index + 1])
    doc_data["consumer_requirements_analysis_questionnaire"] = [consumer_requirements_analysis1,
                                                                consumer_requirements_analysis2]

    product_comparison = process_product_comparison(doc.tables[product_comparison_types_table_index],
                                                    product_recommendation)
    doc_data["product_comparison"] = product_comparison
    return doc_data


if __name__ == '__main__':
    filepath = '../../references/Delaney & Ralley - Preliminary Assessment.docx'
    # with open("../document_generation/result_preliminary_assessment.json", "w") as json_file:
    #     json.dump(parser_doc(filepath), json_file, indent=4)
    print(parser_doc(filepath))
